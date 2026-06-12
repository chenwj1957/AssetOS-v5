from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.tools.build_docx import renderer
from src.tools.build_docx.schema.blocks import (
    CellValue,
    DocumentBlock,
    HorizontalRuleBlock,
    KeyValueTableBlock,
    PageBreakBlock,
    ParagraphBlock,
    SpacerBlock,
    TableBlock,
)
from src.tools.build_docx.schema.docx_spec import DocxSpec
from src.tools.build_docx.validators import validate_docx_spec, validate_output_docx_path


class DocxBuilder:
    def __init__(self) -> None:
        self.document: Document | None = None
        self.spec: DocxSpec | None = None

    def build(self, spec: DocxSpec, output_path: Path) -> Path:
        validate_docx_spec(spec)
        validate_output_docx_path(output_path)
        self.spec = spec
        self.document = Document()
        renderer.apply_page_settings(
            self.document,
            spec.page.size,
            spec.page.orientation,
            spec.page.margins,
        )
        self._apply_normal_style()
        for block in spec.blocks:
            self.render_block(block)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(output_path)
        return output_path

    def render_block(self, block: DocumentBlock, container=None) -> None:
        target = container or self._document()
        if isinstance(block, ParagraphBlock):
            self._render_paragraph(target, block)
        elif isinstance(block, TableBlock):
            self._render_table(target, block)
        elif isinstance(block, KeyValueTableBlock):
            self._render_key_value_table(target, block)
        elif isinstance(block, SpacerBlock):
            self._render_spacer(target, block)
        elif isinstance(block, HorizontalRuleBlock):
            paragraph = target.add_paragraph()
            renderer.add_horizontal_rule(paragraph)
        elif isinstance(block, PageBreakBlock):
            self._document().add_page_break()
        else:
            raise ValueError(f"Unsupported DOCX block: {block}")

    def _render_paragraph(self, target, block: ParagraphBlock) -> None:
        paragraph = target.add_paragraph()
        renderer.apply_paragraph_spacing(paragraph)
        style = self._style(block.style)
        renderer.set_paragraph_alignment(paragraph, block.alignment or (style.alignment if style else None))
        run = paragraph.add_run(block.text)
        renderer.apply_text_style(run, style)

    def _render_spacer(self, target, block: SpacerBlock) -> None:
        paragraph = target.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(block.height)

    def _render_table(self, target, block: TableBlock) -> None:
        row_count = len(block.rows) + (1 if block.headers else 0)
        table = target.add_table(rows=max(row_count, 1), cols=block.columns)
        renderer.align_table(table)
        renderer.set_table_borders(table, block.border)
        renderer.set_column_widths(table, block.column_widths)
        row_index = 0
        if block.headers:
            self._fill_row(table.rows[row_index].cells, block.headers, header=True)
            row_index += 1
        for row in block.rows:
            self._fill_row(table.rows[row_index].cells, row)
            row_index += 1

    def _render_key_value_table(self, target, block: KeyValueTableBlock) -> None:
        table = target.add_table(rows=max(len(block.rows), 1), cols=2)
        renderer.align_table(table)
        renderer.set_table_borders(table, block.border)
        renderer.set_column_widths(table, block.column_widths)
        for index, row in enumerate(block.rows):
            self._fill_row(table.rows[index].cells, row[:2])

    def _fill_row(self, cells, values: list[CellValue], header: bool = False) -> None:
        for index, value in enumerate(values[: len(cells)]):
            style_name = "table_header" if header else None
            if isinstance(value, list):
                renderer.set_cell_blocks(cells[index], value, self)
            else:
                renderer.set_cell_text(cells[index], value, self._style(style_name))

    def _apply_normal_style(self) -> None:
        normal = self._document().styles["Normal"]
        style = self._style("normal")
        if style is None:
            return
        normal.font.name = style.font_name
        normal.font.size = Pt(style.font_size)

    def _style(self, name: str | None):
        if self.spec is None or name is None:
            return None
        return self.spec.styles.get(name)

    def _document(self) -> Document:
        if self.document is None:
            raise RuntimeError("DocxBuilder has not been initialized.")
        return self.document
